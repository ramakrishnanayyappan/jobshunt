from __future__ import annotations

import platform
import subprocess
from pathlib import Path
from typing import List, Optional, Tuple

_PDF_SECTION = {
    "SUMMARY": "PROFESSIONAL SUMMARY",
    "CORE COMPETENCIES": "CORE COMPETENCIES",
    "EXPERIENCE": "PROFESSIONAL EXPERIENCE",
    "EDUCATION": "EDUCATION",
    "CERTIFICATIONS & TRAINING": "CERTIFICATIONS AND TRAINING",
}


def _normalize_header(line: str) -> Optional[str]:
    s = line.strip()
    if s in {
        "SUMMARY",
        "CORE COMPETENCIES",
        "EXPERIENCE",
        "EDUCATION",
        "CERTIFICATIONS & TRAINING",
    }:
        return s
    return None


def parse_txt_content(text: str) -> Tuple[str, str, List[Tuple[str, List[str]]]]:
    lines = text.replace("\r\n", "\n").replace("\r", "\n").split("\n")
    if len(lines) < 4:
        raise ValueError("Résumé text too short: need name, contact, blank line, and sections.")
    name = lines[0].strip()
    contact = lines[1].strip()
    idx = 2
    if idx < len(lines) and not lines[idx].strip():
        idx += 1
    sections: List[Tuple[str, List[str]]] = []
    while idx < len(lines):
        hdr = _normalize_header(lines[idx])
        if hdr is None:
            idx += 1
            continue
        idx += 1
        body: List[str] = []
        while idx < len(lines):
            nxt = _normalize_header(lines[idx])
            if nxt is not None:
                break
            body.append(lines[idx])
            idx += 1
        while body and not body[-1].strip():
            body.pop()
        while body and not body[0].strip():
            body.pop(0)
        sections.append((hdr, body))
    return name, contact, sections


def _escape_rl(text: str) -> str:
    return text.replace("&", "&amp;")


def _build_pdf_impl(name: str, contact: str, sections: List[Tuple[str, List[str]]], path_pdf: Path) -> None:
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_LEFT
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import HRFlowable, Paragraph, SimpleDocTemplate

    base = getSampleStyleSheet()
    styles: dict = {}

    styles["name"] = ParagraphStyle(
        "Name",
        parent=base["Normal"],
        fontName="Helvetica-Bold",
        fontSize=20,
        leading=24,
        alignment=TA_LEFT,
        spaceAfter=6,
        textColor=colors.HexColor("#111111"),
    )
    styles["contact"] = ParagraphStyle(
        "Contact",
        parent=base["Normal"],
        fontName="Helvetica",
        fontSize=9.5,
        leading=12,
        spaceAfter=14,
        textColor=colors.HexColor("#333333"),
    )
    styles["section"] = ParagraphStyle(
        "Section",
        parent=base["Normal"],
        fontName="Helvetica-Bold",
        fontSize=11,
        leading=13,
        spaceBefore=16,
        spaceAfter=8,
        textColor=colors.HexColor("#111111"),
    )
    styles["section_first"] = ParagraphStyle(
        "SectionFirst",
        parent=styles["section"],
        spaceBefore=4,
    )
    styles["body"] = ParagraphStyle(
        "Body",
        parent=base["Normal"],
        fontName="Helvetica",
        fontSize=10.5,
        leading=14,
        spaceAfter=4,
        textColor=colors.HexColor("#222222"),
    )
    styles["skills"] = ParagraphStyle(
        "Skills",
        parent=base["Normal"],
        fontName="Helvetica",
        fontSize=9.5,
        leading=13,
        spaceAfter=4,
        textColor=colors.HexColor("#222222"),
    )
    styles["job_header"] = ParagraphStyle(
        "JobHeader",
        parent=base["Normal"],
        fontName="Helvetica-Bold",
        fontSize=10.5,
        leading=13,
        spaceBefore=12,
        spaceAfter=6,
        textColor=colors.HexColor("#111111"),
    )
    styles["job_first"] = ParagraphStyle(
        "JobFirst",
        parent=styles["job_header"],
        spaceBefore=2,
    )
    styles["bullet"] = ParagraphStyle(
        "Bullet",
        parent=base["Normal"],
        fontName="Helvetica",
        fontSize=10,
        leading=13.5,
        leftIndent=14,
        firstLineIndent=-10,
        spaceAfter=4,
        textColor=colors.HexColor("#222222"),
    )

    story: list = []
    story.append(Paragraph(_escape_rl(name), styles["name"]))
    story.append(Paragraph(_escape_rl(contact), styles["contact"]))
    story.append(
        HRFlowable(
            width="100%",
            thickness=0.5,
            lineCap="round",
            color=colors.HexColor("#cccccc"),
            spaceAfter=12,
        )
    )

    first_block = True
    for hdr_key, body_lines in sections:
        pdf_hdr = _PDF_SECTION.get(hdr_key, hdr_key)
        story.append(
            Paragraph(
                _escape_rl(pdf_hdr),
                styles["section_first"] if first_block else styles["section"],
            )
        )
        first_block = False

        if hdr_key == "EXPERIENCE":
            first_job = True
            for line in body_lines:
                if not line.strip():
                    continue
                bullet = "\u2022"
                stripped = line.lstrip()
                if stripped.startswith(bullet) or stripped.startswith("-"):
                    text = stripped[1:].lstrip()
                    story.append(Paragraph(f"- {_escape_rl(text)}", styles["bullet"]))
                else:
                    st = styles["job_first"] if first_job else styles["job_header"]
                    story.append(Paragraph(_escape_rl(line.strip()), st))
                    first_job = False
        else:
            style = styles["skills"] if hdr_key == "CORE COMPETENCIES" else styles["body"]
            for line in body_lines:
                if not line.strip():
                    continue
                story.append(Paragraph(_escape_rl(line.strip()), style))

    path_pdf.parent.mkdir(parents=True, exist_ok=True)
    doc = SimpleDocTemplate(
        str(path_pdf),
        pagesize=letter,
        leftMargin=0.65 * inch,
        rightMargin=0.65 * inch,
        topMargin=0.55 * inch,
        bottomMargin=0.55 * inch,
        title=f"{name} Résumé",
        author=name,
        subject="Résumé",
    )
    doc.build(story)

    if platform.system() == "Darwin":
        subprocess.run(["xattr", "-cr", str(path_pdf)], check=False)


def build_pdf(resume_text: str, path_pdf: Path) -> None:
    name, contact, sections = parse_txt_content(resume_text)
    _build_pdf_impl(name, contact, sections, path_pdf)


def optional_reserialized_pdf(path_pdf: Path) -> Optional[Path]:
    """Second PDF via pypdf rewrite (vendor-neutral duplicate). Returns path or None."""
    try:
        from pypdf import PdfReader, PdfWriter
    except ImportError:
        return None
    alt = path_pdf.with_name(path_pdf.stem + "_reserialized.pdf")
    try:
        w = PdfWriter()
        r = PdfReader(str(path_pdf))
        for page in r.pages:
            w.add_page(page)
        with open(alt, "wb") as f:
            w.write(f)
        if platform.system() == "Darwin":
            subprocess.run(["xattr", "-cr", str(alt)], check=False)
        return alt
    except Exception:
        return None


def build_docx(resume_text: str, path_docx: Path) -> None:
    from docx import Document
    from docx.enum.text import WD_LINE_SPACING
    from docx.shared import Pt

    name, contact, sections = parse_txt_content(resume_text)

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    pf = normal.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.space_after = Pt(2)

    t = doc.add_paragraph()
    r = t.add_run(name)
    r.bold = True
    r.font.size = Pt(18)
    r.font.name = "Calibri"

    doc.add_paragraph(contact)

    for title, body_lines in sections:
        h = doc.add_paragraph()
        hr = h.add_run(title)
        hr.bold = True
        hr.font.size = Pt(11)
        hr.font.name = "Calibri"

        if not body_lines:
            continue

        for line in body_lines:
            if not line.strip():
                doc.add_paragraph()
                continue
            bullet = "\u2022"
            stripped = line.lstrip()
            if stripped.startswith(bullet) or stripped.startswith("-"):
                text = stripped[1:].lstrip()
                p = doc.add_paragraph(style="List Bullet")
                p.paragraph_format.space_after = Pt(2)
                p.add_run(text)
            else:
                p = doc.add_paragraph(line.strip())
                p.paragraph_format.space_after = Pt(2)

    path_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path_docx)
    if platform.system() == "Darwin":
        subprocess.run(["xattr", "-cr", str(path_docx)], check=False)
